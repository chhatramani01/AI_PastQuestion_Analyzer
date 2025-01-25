import os
from flask import Flask, render_template, request, send_file, redirect, url_for, jsonify
from dotenv import load_dotenv
import google.generativeai as genai
from pdf2image import convert_from_path
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import base64
from io import BytesIO
import logging
from flask import send_from_directory

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['IMAGES_FOLDER'] = 'images'
app.config['OUTPUTS_FOLDER'] = 'outputs'

# Create necessary folders
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['IMAGES_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUTS_FOLDER'], exist_ok=True)

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Gemini API Key
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# Configure Gemini
genai.configure(api_key=GEMINI_API_KEY)
vision_model = genai.GenerativeModel("gemini-1.5-flash")

def upload_pdf(pdf_path):
    """Convert PDF to images and save them to the images/ folder."""
    try:
        images = convert_from_path(pdf_path)
        for i, image in enumerate(images):
            image_path = os.path.join(app.config['IMAGES_FOLDER'], f"page_{i + 1}.png")
            image.save(image_path, "PNG")
            logging.info(f"Saved {image_path}")
        return images
    except Exception as e:
        logging.error(f"Error converting PDF to images: {e}")
        return None

def extract_questions_from_image(image, chapter, topic):
    """Send image to Gemini Vision model and extract questions."""
    buffered = BytesIO()
    image.save(buffered, format="PNG")
    img_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")

    try:
        vision_prompt = (
            f"Extract all questions related to chapter '{chapter}' and topic '{topic}'. "
            "For each question, include the marks and year (if available). "
            "Format the output as follows:\n"
            "**Question:** [Question text] [Marks] [Year]\n"
            "Example:\n"
            "**Question 1:** What is an Instrumentation system? What are the types of errors in instrument explain. [6 marks] [2081 BS]"
        )
        response = vision_model.generate_content(
            [
                vision_prompt,
                {"mime_type": "image/png", "data": img_base64}
            ]
        )
        return response.text
    except Exception as e:
        logging.error(f"Error extracting questions: {e}")
        return None

def save_to_docx(questions, output_path):
    """Save extracted questions to a .docx file."""
    try:
        doc = Document()
        for question in questions:
            doc.add_paragraph(question)
        doc.save(output_path)
        logging.info(f"Questions saved to {output_path}")
    except Exception as e:
        logging.error(f"Error saving to DOCX: {e}")

def analyze_questions(docx_path):
    """Send .docx file to Gemini LLM for analysis."""
    try:
        doc = Document(docx_path)
        content = "\n".join([para.text for para in doc.paragraphs])
        
        llm_prompt = (
            "Analyze the questions in this document. Organize them by topic, count repetitions, and assign importance scores (out of 10). "
            "Also, provide a list of important topics. Format the output as follows:\n"
            "**Topic:** [Topic Name]\n"
            "**Questions:**\n"
            "[Question no] [Question text] [Marks] [Year] (Repeats: [Count]) (Score: [Importance Score])\n"
            "**Important Topics:**\n"
            " [Topic 1]\n"
            " [Topic 2]\n"
            "Example:\n"
            "Topic: Analog-Digital Conversion\n"
            "**Questions:**\n"
            "Q1. What is an Instrumentation system? What are the types of errors in instrument explain. [6 marks] [2081] (Repeats: 4) (Score: 8)\n"
            "**Important Topics:**\n"
            "-> Analog-Digital Conversion\n"
            "-> Instrumentation Errors"
        )
        response = vision_model.generate_content(llm_prompt + "\n\n" + content)
        return response.text
    except Exception as e:
        logging.error(f"Error analyzing questions: {e}")
        return None

def save_to_pdf(analyzed_questions, output_path):
    """Save analyzed questions to a .pdf file with proper formatting and page size."""
    try:
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # Add a title
        title = Paragraph("Analyzed Questions", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))

        # Add analyzed questions
        for line in analyzed_questions.split("\n"):
            if line.strip().startswith("**"):  # Bold headings
                story.append(Paragraph(line.strip(), styles['Heading2']))
            else:  # Normal text
                story.append(Paragraph(line.strip(), styles['Normal']))
            story.append(Spacer(1, 6))

        # Build the PDF
        doc.build(story)
        logging.info(f"Analyzed questions saved to {output_path}")
    except Exception as e:
        logging.error(f"Error saving to PDF: {e}")


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        # Get uploaded file
        file = request.files["file"]
        chapter = request.form["chapter"]
        topic = request.form["topic"]

        if file:
            # Save the uploaded file
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(pdf_path)

            # Step 1: Convert PDF to images
            images = upload_pdf(pdf_path)
            if not images:
                return jsonify({"error": "Error converting PDF to images."})

            # Step 2: Extract questions from images
            all_questions = []
            for image in images:
                questions = extract_questions_from_image(image, chapter, topic)
                if questions:
                    all_questions.extend(questions.split("\n"))

            # Step 3: Save questions to .docx
            docx_path = os.path.join(app.config['OUTPUTS_FOLDER'], "extracted_questions.docx")
            save_to_docx(all_questions, docx_path)

            # Step 4: Analyze questions using Gemini
            analyzed_questions = analyze_questions(docx_path)
            if not analyzed_questions:
                return jsonify({"error": "Error analyzing questions."})

            # Step 5: Save analyzed questions to .pdf
            pdf_path = os.path.join(app.config['OUTPUTS_FOLDER'], "analyzed_questions.pdf")
            save_to_pdf(analyzed_questions, pdf_path)

            # Return JSON response
            return jsonify({
                "analyzed_questions": analyzed_questions,
                "pdf_path": f"/outputs/analyzed_questions.pdf"
            })

    return render_template("index.html")

@app.route("/outputs/<filename>")
def outputs(filename):
    return send_from_directory(app.config['OUTPUTS_FOLDER'], filename)

@app.route("/download")
def download():
    """Download the analyzed questions PDF."""
    pdf_path = os.path.join(app.config['OUTPUTS_FOLDER'], "analyzed_questions.pdf")
    return send_file(pdf_path, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)